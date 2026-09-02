"""Build the editable Word version of Table 2."""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
OUTPUT = RESULTS / "Table2_documented_structure_defense.docx"


INK = "222222"
MID = "555555"
HEADER_FILL = "F2F4F7"
SECTION_FILL = "E8EEF5"
GRID = "B8BDC5"
CONTENT_WIDTH = 9360
CELL_MARGINS = {"top": 55, "bottom": 55, "start": 105, "end": 105}


def set_font(run, size=9, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), GRID)


def apply_table_geometry(table, widths, table_width_dxa=None, indent_dxa=0,
                         cell_margins_dxa=None):
    """Apply fixed Word table widths without external helper modules."""
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            if cell_margins_dxa:
                margins = tc_pr.find(qn("w:tcMar"))
                if margins is None:
                    margins = OxmlElement("w:tcMar")
                    tc_pr.append(margins)
                for side, value in cell_margins_dxa.items():
                    node = margins.find(qn(f"w:{side}"))
                    if node is None:
                        node = OxmlElement(f"w:{side}")
                        margins.append(node)
                    node.set(qn("w:w"), str(value))
                    node.set(qn("w:type"), "dxa")


def section_heading(document, label):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    set_font(paragraph.add_run(label), size=9.5, bold=True)


def add_table(document, headers, rows, widths, alignments=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, HEADER_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
        )
        paragraph.paragraph_format.space_after = Pt(0)
        set_font(paragraph.add_run(header), size=8.6, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            alignment = alignments[index] if alignments else (
                WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.alignment = alignment
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            set_font(paragraph.add_run(str(value)), size=8.6)
    apply_table_geometry(
        table, widths, table_width_dxa=CONTENT_WIDTH, indent_dxa=105,
        cell_margins_dxa=CELL_MARGINS,
    )
    set_borders(table)
    return table


def fmt_p(value):
    if value < .001:
        return "<0.001"
    return f"{value:.3f}"


def main():
    direct = pd.read_csv(RESULTS / "table2a_defense_direct_effect.csv")
    outcomes = pd.read_csv(RESULTS / "table2b_defense_outcomes.csv")
    spillover = pd.read_csv(RESULTS / "table2c_defense_spillover.csv")
    mechanism = pd.read_csv(RESULTS / "table2d_defense_mechanism.csv")

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    set_font(
        title.add_run("Table 2 | Documented structure defense"),
        size=11.5, bold=True,
    )

    section_heading(document, "A. Matched association with focal-structure survival")
    direct_rows = []
    for row in direct.itertuples(index=False):
        direct_rows.append([
            row.sample, f"{int(row.matched_pairs):,}",
            f"{row.survival_difference_pp:.1f} "
            f"({row.ci_lo:.1f} to {row.ci_hi:.1f})",
            fmt_p(row.p_value),
        ])
    add_table(
        document,
        ["Sample", "Matched pairs", "Difference, pp (95% CI)", "P value"],
        direct_rows, [1870, 1510, 4250, 1730],
    )

    section_heading(document, "B. Outcomes in the pooled matched sample")
    outcome_rows = [[
        row.group, f"{row.Undamaged:.1f}", f"{row.Partial:.1f}",
        f"{row.Destroyed:.1f}",
    ] for row in outcomes.itertuples(index=False)]
    add_table(
        document,
        ["Group", "Undamaged, %", "Partial, %", "Destroyed, %"],
        outcome_rows, [2520, 2280, 2280, 2280],
    )

    section_heading(document, "C. Directional spillover in the restricted, rematched sample")
    pooled = spillover[spillover["sample"].eq("Pooled")].iloc[0]
    spillover_rows = [
        ["First-stage focal survival", f"{100 * pooled.first_stage:.1f}"],
        [
            "Placebo-corrected directional contrast",
            f"{100 * pooled.directional_contrast:.1f} "
            f"({100 * pooled.directional_contrast_lo:.1f} to "
            f"{100 * pooled.directional_contrast_hi:.1f})",
        ],
        [
            "Local IV estimate per focal structure saved",
            f"{100 * pooled.local_iv:.1f} "
            f"({100 * pooled.local_iv_lo:.1f} to "
            f"{100 * pooled.local_iv_hi:.1f})",
        ],
    ]
    add_table(
        document, ["Estimand", "Estimate, pp (95% CI)"],
        spillover_rows, [5650, 3710],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )

    section_heading(document, "D. Emitter-removal mechanism")
    mechanism_rows = []
    for row in mechanism.itertuples(index=False):
        mechanism_rows.append([
            row.model, row.term, f"{row.ame_pp_per_sd:.2f}",
            fmt_p(row.p_value),
        ])
    add_table(
        document,
        ["Model", "Term", "AME, pp per s.d.", "P value"],
        mechanism_rows, [2500, 3870, 1850, 1140],
    )

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.keep_together = True
    note_run = note.add_run(
        "Notes: Positive direct differences indicate greater survival with "
        "documented defense; negative spillover estimates indicate fewer "
        "destroyed down-fire neighbors. The spillover analysis includes 225 "
        "rematched pairs (450 structures) with at least 75% of up-fire neighbors "
        "destroyed and observed undefended neighbors within 100 ft in both "
        "directions. AME, average marginal effect; IV, instrumental variable; "
        "pp, percentage points; s.d., standard deviation."
    )
    set_font(note_run, size=7.8, color=MID)

    document.core_properties.title = "Table 2 | Documented structure defense"
    document.core_properties.subject = "2025 Los Angeles fire analysis"
    document.core_properties.author = ""
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
